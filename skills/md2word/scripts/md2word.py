#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换工具（完整版）
支持自动格式化为法律文书标准格式，现已支持表格转换和格式保留：
- 页面大小：A4 (21cm × 29.7cm)
- 字体：仿宋_GB2312，黑色
- 字号：小四 (12pt)
- 行距：1.5倍
- 首行缩进：2个字符
- 一级标题：小三号，居中加粗，段前段后0.5行
- 其他内容：两端对齐
- 页边距：上下2.54cm，左右3.18cm
- 引号转换：自动将英文引号转换为中文引号
- 页码设置：自动添加页脚页码（格式：1/x，Times New Roman五号）
- 表格转换：支持Markdown表格转换为Word表格，自动设置边框和格式
- 格式支持：支持**加粗**、*斜体*、<u>下划线</u>、~~删除线~~等格式

使用方法：
1. 简单使用：将此脚本放在包含.md文件的文件夹中，运行脚本
2. 指定文件：python md_to_word_converter_complete.py input.md output.docx
3. 使用模板：python md_to_word_converter_complete.py input.md output.docx template.docx
4. 自动模板：程序会自动查找同目录下的.docx文件作为模板（优先使用包含'模板'或'template'的文件）
"""

import os
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re
import glob
import requests
import base64
import io
from PIL import Image
import tempfile
import subprocess
import shutil
import time
from bs4 import BeautifulSoup

# 导入配置模块
from config import Config, load_config, get_preset, get_default_preset, merge_configs, list_presets

# 全局配置对象
_current_config: Config = None

def get_config() -> Config:
    """获取当前配置"""
    global _current_config
    if _current_config is None:
        _current_config = get_default_preset()
    return _current_config

def set_config(config: Config):
    """设置当前配置"""
    global _current_config
    _current_config = config

# 全局图片参数（可按需调整）
# 图片显示大小（更大更清晰）：
# - 将默认占比由 70% 提升到 92%
# - 将最大宽度由 10.5cm 提升到 14.2cm（A4 可用宽约 14.64cm）
IMAGE_DISPLAY_RATIO = 0.92         # 相对于页面可用宽度的比例
IMAGE_MAX_DISPLAY_CM = 14.2        # 图片最大显示宽度（cm）
IMAGE_TARGET_DPI = 260             # 目标DPI（用于下采样像素宽度计算）

def get_image_output_path(md_file_path, png_filename):
    """获取图片输出路径，确保目录存在"""
    md_dir = os.path.dirname(os.path.abspath(md_file_path))
    # 基于Markdown文件名创建子目录
    md_filename_base = os.path.splitext(os.path.basename(md_file_path))[0]
    image_dir = os.path.join(md_dir, f"{md_filename_base}_images")
    
    if not os.path.exists(image_dir):
        try:
            os.makedirs(image_dir)
            print(f"📂 创建图片目录: {os.path.relpath(image_dir)}")
        except OSError as e:
            print(f"⚠️ 创建目录失败: {e}")
            return None
            
    return os.path.join(image_dir, png_filename)


def create_word_document(md_file_path, output_path, template_file=None, config: Config = None):
    """
    从Markdown文件创建格式化的Word文档

    Args:
        md_file_path: Markdown文件路径
        output_path: 输出Word文件路径
        template_file: Word模板文件路径（可选）
        config: 配置对象（可选）
    """
    if config is None:
        config = get_config()

    print(f"📄 正在处理: {md_file_path}")
    print(f"📋 使用配置: {config.name}")

    # 添加引号调试
    if config.get('quotes.convert_to_chinese', True):
        debug_quotes_in_file(md_file_path)

    # 创建或加载文档
    if template_file and template_file != "none" and os.path.exists(template_file):
        print(f"📋 使用模板文件: {os.path.basename(template_file)}")
        doc = Document(template_file)
        # 清空模板内容
        try:
            # 清空段落和表格
            for paragraph in list(doc.paragraphs):
                if paragraph != doc.paragraphs[0]:
                    p = paragraph._element
                    p.getparent().remove(p)
                else:
                    paragraph.clear()

            for table in list(doc.tables):
                t = table._element
                t.getparent().remove(t)
        except Exception as e:
            print(f"⚠️ 清空模板内容失败: {e}")
    else:
        print("📄 创建新文档（不使用模板）")
        doc = Document()

    # 设置默认字体
    try:
        normal_style = doc.styles['Normal']
        font_config = config.get('fonts.default', {})
        normal_style.font.name = font_config.get('ascii', 'Times New Roman')
        normal_style.font.size = Pt(font_config.get('size', 10.5))
        normal_style._element.rPr.rFonts.set(qn('w:ascii'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('name', '仿宋_GB2312'))
        normal_style._element.rPr.rFonts.set(qn('w:cs'), font_config.get('ascii', 'Times New Roman'))
    except Exception as _:
        pass

    # 设置页面大小和页边距
    sections = doc.sections
    for section in sections:
        page_config = config.get('page', {})
        section.page_width = Cm(page_config.get('width', 21.0))
        section.page_height = Cm(page_config.get('height', 29.7))
        section.top_margin = Cm(page_config.get('margin_top', 2.54))
        section.bottom_margin = Cm(page_config.get('margin_bottom', 2.54))
        section.left_margin = Cm(page_config.get('margin_left', 3.18))
        section.right_margin = Cm(page_config.get('margin_right', 3.18))
    
    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # 如果UTF-8失败，尝试其他编码
        with open(md_file_path, 'r', encoding='gbk') as f:
            content = f.read()
    
    # 按行处理内容，保留原始行结构
    lines = content.split('\n')
    processed_lines = lines  # 保留所有行，包括空行
    
    # 处理表格和图表
    has_body_before_first_h2 = False
    has_seen_h2 = False
    i = 0
    while i < len(processed_lines):
        line = processed_lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 检查是否是Mermaid图表开始（兼容 ``` mermaid / ```mermaid ）
        if re.match(r'^```\s*mermaid\b', line):
            mermaid_lines = []
            i += 1  # 跳过开始标记
            
            # 收集Mermaid代码
            while i < len(processed_lines) and not processed_lines[i].strip().startswith('```'):
                mermaid_lines.append(processed_lines[i])
                i += 1
            
            if i < len(processed_lines):
                i += 1  # 跳过结束标记
            
            # 处理Mermaid图表
            if mermaid_lines:
                mermaid_code = '\n'.join(mermaid_lines)
                create_mermaid_chart(doc, mermaid_code, md_file_path)
                # 插入了实质内容
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Mermaid图表")
            continue
        
        # 处理代码块 ```lang ... ```
        if line.startswith('```'):
            code_lines = []
            language = line[3:].strip()
            i += 1
            while i < len(processed_lines) and not processed_lines[i].strip().startswith('```'):
                code_lines.append(processed_lines[i])
                i += 1
            if i < len(processed_lines):
                i += 1  # 跳过结束标记
            add_code_block(doc, code_lines, language)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            print("✅ 处理代码块")
            continue
        
        # 检查是否是HTML表格开始
        if '<table>' in line.lower():
            html_table_content = []
            table_start = i
            
            # 收集HTML表格的所有内容
            while i < len(processed_lines):
                current_line = processed_lines[i]
                html_table_content.append(current_line)
                if '</table>' in current_line.lower():
                    i += 1
                    break
                i += 1
            
            # 处理HTML表格
            if html_table_content:
                html_content = '\n'.join(html_table_content)
                create_word_table_from_html(doc, html_content)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
        
        # 检查是否是Markdown表格开始
        if is_table_row(line):
            table_lines = []
            table_start = i
            
            # 收集表格的所有行
            while i < len(processed_lines) and is_table_row(processed_lines[i].strip()):
                table_lines.append(processed_lines[i].strip())
                i += 1
            
            # 处理表格
            if len(table_lines) >= 2:  # 至少要有标题行和分隔行
                create_word_table(doc, table_lines)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Markdown表格: {len(table_lines)} 行")
            continue
        
        # 分割线
        if line in ['---', '***', '___']:
            add_horizontal_line(doc)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 任务列表
        if line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            add_task_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 无序列表
        if line.startswith(('- ', '* ', '+ ')):
            add_bullet_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+\.\s', line):
            add_numbered_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 引用块（处理多行引用）
        if line.startswith('>'):
            quote_lines = []
            # 收集连续的引用行
            while i < len(lines) and lines[i].startswith('>'):
                quote_content = lines[i][1:].strip()  # 移除 > 符号
                quote_lines.append(quote_content)  # 添加所有行，包括空行
                i += 1
            
            # 将多行引用合并为一个引用块
            if quote_lines:
                full_quote = '\n'.join(quote_lines)
                add_quote(doc, full_quote)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue
            
        # 判断标题级别
        if line.startswith('# '):
            # 一级标题：小三号，居中加粗，段前段后0.5行
            title = line[2:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=1)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=1)
            
        elif line.startswith('## '):
            # 二级标题：加粗，支持内部格式
            # 在条件满足时，标题前插入一个空行
            if has_seen_h2 or has_body_before_first_h2:
                doc.add_paragraph("")
            title = line[3:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=2)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=2)
            has_seen_h2 = True
            
        elif line.startswith('### '):
            # 三级标题：不加粗，但支持内部格式
            title = line[4:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=3)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=3)
            
        elif line.startswith('#### '):
            # 四级标题：不加粗，但支持内部格式
            title = line[5:].strip()
            title = convert_quotes_to_chinese(title)  # 转换引号
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=4)  # 使用格式解析处理标题内容
            set_paragraph_format(p, title_level=4)
            
        else:
            # 正文段落
            if line:
                p = doc.add_paragraph()
                parse_text_formatting(p, line)
                set_paragraph_format(p)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
        
        i += 1
    
    # 添加页码
    add_page_number(doc)
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ Word文档已生成: {output_path}")

def preprocess_mermaid_code(mermaid_code: str) -> str:
    """预处理Mermaid源码，避免Mermaid v11 对标签内Markdown解析导致的
    "Unsupported markdown: list"/"codespan" 等错误。
    - 将行首的 "- ", "* " 项目符号替换为 "• "（兜底）
    - 将编号列表的 "1. " 改为 "1: "（兜底）
    - 将反引号 ` 替换为普通单引号 '，以避免 codespan 报错
    - 重点：对节点标签内部（[...], (...), ((...)), {...}, >...], ["..."], ("...")) 的起始列表标记进行替换
    该处理为无害替换，不影响边、样式等语句。
    """
    import re

    s = mermaid_code

    # 反引号替换，避免 codespan 被解析
    s = s.replace("`", "'")

    # 1) 针对节点标签内部：有序列表 1. -> 1:
    def _repl_number_dot(m: re.Match) -> str:
        brace = m.group('brace')
        quote = m.group('quote') or ''
        num = m.group('num')
        return f"{brace}{quote}{num}: "

    s = re.sub(r"(?m)(?P<brace>[\[\({\>])(?P<quote>\"?\s*)(?P<num>\d+)\.\s", _repl_number_dot, s)

    # 2) 针对节点标签内部：无序列表 - / * -> •
    def _repl_bullet(m: re.Match) -> str:
        brace = m.group('brace')
        quote = m.group('quote') or ''
        return f"{brace}{quote}• "

    s = re.sub(r"(?m)(?P<brace>[\[\({\>])(?P<quote>\"?\s*)[-*]\s", _repl_bullet, s)

    # 3) 兜底：整行以列表开头的情况（极少出现在Mermaid内，但保留以防万一）
    s = re.sub(r"(?m)^(\s*)-\s+", r"\1• ", s)
    s = re.sub(r"(?m)^(\s*)\*\s+", r"\1• ", s)
    s = re.sub(r"(?m)^(\s*)(\d+)\.\s+", r"\1\2: ", s)

    return s

def create_mermaid_chart(doc, mermaid_code, md_file_path):
    """将Mermaid图表转换为图片并插入Word文档（本地渲染优先）"""

    # 预处理，规避 Mermaid 11 对列表/反引号的 Markdown 解析造成的报错
    mermaid_code = preprocess_mermaid_code(mermaid_code)

    # 首先尝试本地渲染
    local_success = try_local_mermaid_render(doc, mermaid_code, md_file_path)
    if local_success:
        return

    # 仅使用本地渲染：失败则改为文本，不再尝试在线服务
    print("⚠️ 本地渲染失败，已禁用在线服务，使用文本替代")
    create_fallback_text(doc, mermaid_code)

def try_local_mermaid_render(doc, mermaid_code, md_file_path):
    """尝试使用本地mermaid-cli渲染图表"""
    
    # 为Mermaid文件和输出图片准备路径
    timestamp = str(int(time.time() * 1000))
    mmd_filename = f"mermaid-src-{timestamp}.mmd"
    png_filename = f"mermaid-chart-{timestamp}.png"
    
    # 获取保存图片的最终路径
    output_png_path = get_image_output_path(md_file_path, png_filename)
    if not output_png_path:
        print("⚠️ 无法获取图片输出路径，跳过本地渲染。")
        return False
        
    # 临时文件放在脚本所在目录，避免 cwd 不一致导致路径问题
    script_dir = os.path.dirname(os.path.abspath(__file__))
    temp_mmd_path = os.path.join(script_dir, mmd_filename)

    try:
        print("🖥️ 尝试本地Mermaid渲染...")
        
        # 创建临时的.mmd文件
        with open(temp_mmd_path, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)
        
        # 检查 mmdc 命令：优先环境变量 MMDCCMD，其次脚本同目录 node_modules，再其次系统 PATH
        mmdc_env = os.environ.get('MMDCCMD', '').strip()
        mmdc_path = mmdc_env if mmdc_env else os.path.join(script_dir, "node_modules", ".bin", "mmdc")
        if not os.path.exists(mmdc_path):
            mmdc_path = shutil.which("mmdc") or ""
        if not mmdc_path:
            print("⚠️ 本地 mmdc 命令未找到（已跳过本地渲染）")
            return False
        
        # 使用mmdc命令生成高分辨率PNG图片
        # 绝对路径，配置文件若存在则使用
        abs_in = os.path.abspath(temp_mmd_path)
        abs_out = os.path.abspath(output_png_path)
        cfg = os.path.join(script_dir, "mermaid-config.json")
        cmd = [mmdc_path, "-i", abs_in, "-o", abs_out, "-t", "neutral", "-w", "2200", "-H", "1500", "--scale", "2.0"]
        if os.path.exists(cfg):
            cmd.extend(["-c", cfg])
        
        print(f"🔧 执行命令: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            print(f"⚠️ mmdc 命令执行失败: {result.stderr}")
            return False
        
        # 检查生成的PNG文件是否存在
        if not os.path.exists(output_png_path):
            print("⚠️ PNG文件未生成")
            return False
        
        # 加载图片并插入Word
        image = Image.open(output_png_path)
        insert_image_to_word(doc, image)
        
        print(f"✅ 本地Mermaid图表渲染成功！图片已保存至: {os.path.relpath(output_png_path)}")
        return True
        
    except subprocess.TimeoutExpired:
        print("⚠️ mmdc命令执行超时")
        return False
    except Exception as e:
        print(f"⚠️ 本地渲染失败: {e}")
        return False
    finally:
        # 无论成功与否，都清理临时的mmd文件
        if os.path.exists(temp_mmd_path):
            try:
                os.unlink(temp_mmd_path)
            except:
                pass

def _postprocess_image_for_word(image, target_display_cm, target_dpi=IMAGE_TARGET_DPI):
    """根据目标显示宽度与DPI对图像进行高质量下采样，控制体积并保持清晰度"""
    try:
        # 目标像素宽度 = 目标显示英寸 * 目标DPI
        target_inches = float(target_display_cm) / 2.54
        target_px_width = max(1, int(target_inches * target_dpi))
        if image.width > target_px_width:
            new_height = int(image.height * (target_px_width / image.width))
            image = image.resize((target_px_width, new_height), Image.LANCZOS)
    except Exception:
        pass
    return image

def insert_image_to_word(doc, image):
    """将PIL图片对象插入Word文档"""
    config = get_config()
    image_config = config.get('image', {})
    page_config = config.get('page', {})

    display_ratio = image_config.get('display_ratio', 0.92)
    max_width_cm = image_config.get('max_width_cm', 14.2)
    target_dpi = image_config.get('target_dpi', 260)

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        # 计算页面可用宽度与目标插入宽度（限制较小以保持版面协调）
        page_width = page_config.get('width', 21.0)
        margin_left = page_config.get('margin_left', 3.18)
        margin_right = page_config.get('margin_right', 3.18)
        available_width_cm = page_width - margin_left - margin_right
        target_display_cm = min(available_width_cm * display_ratio, max_width_cm)
        # 按目标DPI对图像进行高质量下采样（清晰但不臃肿）
        image = _postprocess_image_for_word(image, target_display_cm, target_dpi=target_dpi)
        # 使用高压缩PNG保存，进一步降低体积
        try:
            image.save(temp_file.name, format='PNG', optimize=True, compress_level=9)
        except Exception:
            image.save(temp_file.name, format='PNG', optimize=True)
        temp_filename = temp_file.name

    try:
        # 在Word中插入图片
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 统一的目标显示宽度，保持版面一致性
        img_width_cm = target_display_cm

        run = paragraph.add_run()
        run.add_picture(temp_filename, width=Cm(img_width_cm))

    finally:
        # 删除临时文件
        try:
            os.unlink(temp_filename)
        except:
            pass

def create_fallback_text(doc, mermaid_code):
    """创建后备文本方案"""
    # 解析图表类型并创建简化版本
    if 'graph' in mermaid_code.lower():
        create_simple_diagram_text(doc, mermaid_code)
    elif 'pie' in mermaid_code.lower():
        create_simple_pie_text(doc, mermaid_code)
    elif 'gantt' in mermaid_code.lower():
        create_simple_gantt_text(doc, mermaid_code)
    else:
        # 默认处理
        p = doc.add_paragraph()
        run = p.add_run("【图表内容】")
        run.bold = True
        p.add_run("\n" + mermaid_code)
        set_paragraph_format(p)

def create_simple_diagram_text(doc, mermaid_code):
    """创建简化的图表文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【流程图】")
    run.bold = True
    
    # 解析节点和连接关系
    lines = mermaid_code.split('\n')
    nodes = {}
    connections = []
    
    for line in lines:
        line = line.strip()
        if '-->' in line or '->' in line:
            parts = line.split('-->' if '-->' in line else '->')
            if len(parts) == 2:
                from_node = parts[0].strip()
                to_node = parts[1].strip()
                connections.append(f"{from_node} → {to_node}")
        elif '[' in line and ']' in line:
            # 解析节点定义
            match = re.search(r'(\w+)\["([^"]+)"\]', line)
            if match:
                node_id, node_text = match.groups()
                nodes[node_id] = node_text
    
    # 添加解析结果
    if connections:
        p.add_run("\n主要流程:")
        for conn in connections[:8]:  # 最多显示8个连接
            p.add_run(f"\n• {conn}")
    
    set_paragraph_format(p)

def create_simple_pie_text(doc, mermaid_code):
    """创建简化的饼图文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【数据分析】")
    run.bold = True
    
    # 解析饼图数据
    lines = mermaid_code.split('\n')
    for line in lines:
        if ':' in line and '"' in line:
            # 解析数据项
            match = re.search(r'"([^"]+)"\s*:\s*(\d+(?:\.\d+)?)', line)
            if match:
                label, value = match.groups()
                p.add_run(f"\n• {label}: {value}")
    
    set_paragraph_format(p)

def create_simple_gantt_text(doc, mermaid_code):
    """创建简化的甘特图文本描述"""
    p = doc.add_paragraph()
    run = p.add_run("【时间安排】")
    run.bold = True
    
    # 解析甘特图任务
    lines = mermaid_code.split('\n')
    current_section = ""
    
    for line in lines:
        line = line.strip()
        if line.startswith('section '):
            current_section = line.replace('section ', '')
            p.add_run(f"\n\n{current_section}:")
        elif ':' in line and not line.startswith('title'):
            # 解析任务
            task = line.split(':')[0].strip()
            p.add_run(f"\n• {task}")
    
    set_paragraph_format(p)

def is_separator_line(line):
    """判断是否是表格分隔行。分隔行必须包含'-'，且只能包含'|', '-', ':', ' '等符号。"""
    line = line.strip()
    if not line or '-' not in line:
        return False
    return all(c in '|-: 	' for c in line)


def is_table_row(line):
    """判断是否是表格行"""
    if not line or not line.strip():
        return False
    
    line = line.strip()
    
    # 检查是否是分隔行
    if is_separator_line(line):
        return True
    
    # 检查是否是数据行（包含 |）
    # 这里的逻辑保持宽松，依赖于主循环中对其他块级元素的优先判断
    if '|' in line:
        return True
    
    return False

def create_word_table(doc, table_lines):
    """从Markdown表格行创建Word表格"""
    
    if len(table_lines) < 2:
        return
    
    # 解析表格数据
    rows_data = []
    header_row = None
    
    for i, line in enumerate(table_lines):
        # 跳过分隔行（包含横线的行）
        if is_separator_line(line):
            continue
        
        # 解析单元格
        cells = parse_table_row(line)
        if cells:
            if header_row is None:
                header_row = cells
            else:
                rows_data.append(cells)
    
    if not header_row:
        return
    
    # 确定列数
    max_cols = len(header_row)
    for row in rows_data:
        max_cols = max(max_cols, len(row))
    
    # 创建Word表格
    total_rows = 1 + len(rows_data)  # 标题行 + 数据行
    table = doc.add_table(rows=total_rows, cols=max_cols)

    # 获取表格配置
    config = get_config()
    table_config = config.get('table', {})
    border_enabled = table_config.get('border_enabled', True)
    border_color = table_config.get('border_color', '#000000')
    border_width = table_config.get('border_width', 4)
    row_height_cm = table_config.get('row_height_cm', 0.8)
    alignment_str = table_config.get('alignment', 'center')
    line_spacing = table_config.get('line_spacing', 1.2)
    cell_margin = table_config.get('cell_margin', {})
    vertical_align_str = table_config.get('vertical_align', 'center')

    # 设置表格对齐方式
    alignment_map = {
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'right': WD_TABLE_ALIGNMENT.RIGHT
    }
    table.alignment = alignment_map.get(alignment_str.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # 设置垂直对齐
    vertical_align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'center': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    vertical_align = vertical_align_map.get(vertical_align_str.lower(), WD_ALIGN_VERTICAL.CENTER)

    # 统一设置边框和内边距、行高等
    if border_enabled:
        try:
            tbl = table._tbl
            color = border_color.lstrip('#')
            borders_xml = f'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            tbl.tblPr.append(parse_xml(borders_xml))
        except Exception:
            pass

    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = f'''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass

    # 行高与段落行距统一
    try:
        for row in table.rows:
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                cell.vertical_alignment = vertical_align
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = line_spacing
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass
    
    # 填充标题行
    header_cells = table.rows[0].cells
    for j, cell_text in enumerate(header_row):
        if j < len(header_cells):
            cell = header_cells[j]
            # 处理表格单元格中的格式
            if contains_markdown_formatting(cell_text.strip()):
                parse_table_cell_formatting(cell, cell_text.strip(), is_header=True)
            else:
                cell.text = convert_quotes_to_chinese(cell_text.strip())
                set_table_cell_format(cell, is_header=True)
    
    # 填充数据行
    for i, row_data in enumerate(rows_data):
        if i + 1 < len(table.rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    # 处理表格单元格中的格式
                    if contains_markdown_formatting(cell_text.strip()):
                        parse_table_cell_formatting(cell, cell_text.strip(), is_header=False)
                    else:
                        cell.text = convert_quotes_to_chinese(cell_text.strip())
                        set_table_cell_format(cell, is_header=False)
    
    # 调整列宽
    adjust_table_column_width(table)

def parse_table_row(line):
    """解析表格行，提取单元格内容"""
    if not line or not line.strip():
        return []
    
    line = line.strip()
    
    # 移除开头和结尾的 |
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]
    
    # 过滤掉空单元格（但保留有意义的空单元格）
    return cells

def contains_markdown_formatting(text):
    """检查文本是否包含Markdown格式标记"""
    format_patterns = [
        r'\*\*\*.*?\*\*\*',  # 加粗斜体
        r'\*\*.*?\*\*',      # 加粗
        r'\*.*?\*',          # 斜体
        r'___.*?___',        # 加粗斜体
        r'__.*?__',          # 加粗
        r'_.*?_',            # 斜体
        r'<u>.*?</u>',       # 下划线
        r'~~.*?~~',          # 删除线
        r'`.*?`',            # 行内代码
        r'<br\s*/?>',       # 换行标签
        r'\$.*?\$',         # LaTeX数学公式
    ]
    
    for pattern in format_patterns:
        if re.search(pattern, text):
            return True
    return False

def parse_table_cell_formatting(cell, text, is_header=False):
    """解析表格单元格中的格式化文本"""
    # 清空单元格
    cell.text = ""
    
    # 转换引号
    text = convert_quotes_to_chinese(text)
    
    # 支持<br>换行：拆分后逐段处理
    parts_by_br = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    
    # 解析格式
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]
    
    for idx, segment in enumerate(parts_by_br):
        if idx > 0:
            cell.paragraphs[0].add_run().add_break()
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = cell.paragraphs[0].add_run(part_text)
                set_table_run_format(run, formats, is_header)

def set_table_run_format(run, formats, is_header=False):
    """设置表格单元格run格式"""
    config = get_config()

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.size = Pt(font_size)
    font.color.rgb = hex_to_rgb(color_hex)
    font.bold = bold if is_header else False

    # 设置字体映射：英文和数字用Times New Roman，中文用仿宋_GB2312
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

    # 应用Markdown格式
    if formats.get('bold', False):
        font.bold = True
    if formats.get('italic', False):
        font.italic = True
    if formats.get('underline', False):
        font.underline = True
    if formats.get('strikethrough', False):
        font.strike = True
    if formats.get('code', False):
        # 表格中代码使用Times New Roman，稍小字号
        code_config = config.get('inline_code', {})
        font.name = code_config.get('font', 'Times New Roman')
        font.size = Pt(9)
        font.color.rgb = hex_to_rgb(code_config.get('color', '#333333'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return
    if formats.get('math', False):
        # 表格中数学公式使用Times New Roman，斜体，深蓝色
        math_config = config.get('math', {})
        font.name = math_config.get('font', 'Times New Roman')
        font.size = Pt(math_config.get('size', 10))
        font.italic = math_config.get('italic', True)
        font.color.rgb = hex_to_rgb(math_config.get('color', '#00008B'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return

def set_table_cell_format(cell, is_header=False):
    """设置表格单元格格式"""
    config = get_config()
    table_config = config.get('table', {})
    line_spacing = table_config.get('line_spacing', 1.2)

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    # 设置段落格式
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing

        # 设置文字格式
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            font.color.rgb = hex_to_rgb(color_hex)
            font.bold = bold if is_header else False

            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def adjust_table_column_width(table):
    """调整表格列宽"""
    try:
        # 获取表格总宽度（页面宽度减去页边距）
        available_width = Cm(21.0 - 3.18 * 2)  # A4宽度减去左右页边距
        
        # 平均分配列宽
        col_count = len(table.columns)
        if col_count > 0:
            col_width = int(available_width / col_count)  # 转换为整数
            for column in table.columns:
                column.width = col_width
    except Exception as e:
        print(f"⚠️  表格列宽调整失败: {e}")

def parse_html_table(html_content):
    """解析HTML表格内容，返回表格数据"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        table = soup.find('table')
        if not table:
            return None
        
        rows_data = []
        for tr in table.find_all('tr'):
            row_cells = []
            for cell in tr.find_all(['td', 'th']):
                # 获取单元格文本内容，保留基本格式
                cell_text = cell.get_text(strip=True)
                row_cells.append(cell_text)
            if row_cells:  # 只添加非空行
                rows_data.append(row_cells)
        
        return rows_data
    except Exception as e:
        print(f"⚠️  HTML表格解析失败: {e}")
        return None

def create_word_table_from_html(doc, html_content):
    """从HTML表格创建Word表格"""
    rows_data = parse_html_table(html_content)
    if not rows_data or len(rows_data) < 1:
        print("⚠️  HTML表格数据为空或格式不正确")
        return

    # 获取表格配置
    config = get_config()
    table_config = config.get('table', {})
    border_enabled = table_config.get('border_enabled', True)
    border_color = table_config.get('border_color', '#000000')
    border_width = table_config.get('border_width', 4)
    row_height_cm = table_config.get('row_height_cm', 0.8)
    line_spacing = table_config.get('line_spacing', 1.2)
    cell_margin = table_config.get('cell_margin', {})
    vertical_align_str = table_config.get('vertical_align', 'center')

    # 创建Word表格
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))

    # 设置表格对齐方式
    alignment_str = table_config.get('alignment', 'center')
    alignment_map = {
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'right': WD_TABLE_ALIGNMENT.RIGHT
    }
    table.alignment = alignment_map.get(alignment_str.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # 设置垂直对齐
    vertical_align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'center': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    vertical_align = vertical_align_map.get(vertical_align_str.lower(), WD_ALIGN_VERTICAL.CENTER)

    # 设置表格边框和单元格边距
    if border_enabled:
        try:
            tbl = table._tbl
            color = border_color.lstrip('#')
            borders_xml = f'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            tbl.tblPr.append(parse_xml(borders_xml))
        except Exception:
            pass

    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = f'''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass

    # 设置行高和单元格对齐
    try:
        for row in table.rows:
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                cell.vertical_alignment = vertical_align
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = line_spacing
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass

    # 填充表格数据
    for i, row_data in enumerate(rows_data):
        if i < len(table.rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    cell.text = convert_quotes_to_chinese(cell_text.strip())
                    # 第一行作为标题行处理
                    set_table_cell_format(cell, is_header=(i == 0))

    # 调整列宽
    adjust_table_column_width(table)
    print(f"✅ 处理HTML表格: {len(rows_data)} 行")

def convert_quotes_to_chinese(text):
    """将英文引号转换为中文引号（交替状态机版）
    规则：
    - 将直双引号 " 转为中文开/闭引号 " "（交替状态：开→闭→开→闭...）
    - 将直单引号 ' 转为中文开/闭引号 ' '，但保留英文缩写/所有格中的撇号（如 don't, John's）
    - 避免转换代码片段中的引号（由反引号 ` 包裹）
    """
    if not text:
        return text

    original_text = text

    # 若无需要处理的引号，直接返回
    if ('"' not in text) and ("'" not in text):
        return text

    result = []
    i = 0
    in_code = False  # 是否处于 `code` 片段中

    # 交替状态机：0=等待开引号，1=等待闭引号
    double_quote_state = 0
    single_quote_state = 0

    while i < len(text):
        ch = text[i]

        # 处理反引号包裹的代码片段，保持原样
        if ch == '`':
            # 统计连续反引号的数量（支持 ``` 块 及 ` 行内`）
            j = i + 1
            while j < len(text) and text[j] == '`':
                j += 1
            backtick_count = j - i
            result.append('`' * backtick_count)
            in_code = not in_code  # 简化处理：遇到成组反引号时翻转状态
            i = j
            continue

        if in_code:
            # 代码片段内不做引号更换
            result.append(ch)
            i += 1
            continue

        if ch == '"':
            # 使用交替状态机：第一个是开引号，第二个是闭引号，以此类推
            if double_quote_state == 0:
                result.append('\u201c')  # 中文开双引号 "
                double_quote_state = 1  # 下一个是闭引号
            else:
                result.append('\u201d')  # 中文闭双引号 "
                double_quote_state = 0  # 重置，下一个是开引号
            i += 1
            continue

        if ch == "'":
            # 保留英文缩写/所有格中的撇号：字母-撇号-字母
            prev_c = text[i - 1] if i > 0 else ''
            next_c = text[i + 1] if i + 1 < len(text) else ''
            if prev_c.isalpha() and next_c.isalpha():
                result.append("'")
                i += 1
                continue

            # 使用交替状态机
            if single_quote_state == 0:
                result.append('\u2018')  # 中文开单引号 '
                single_quote_state = 1
            else:
                result.append('\u2019')  # 中文闭单引号 '
                single_quote_state = 0
            i += 1
            continue

        # 其它字符保持
        result.append(ch)
        i += 1

    text = ''.join(result)

    if text != original_text:
        print(f"✅ 引号转换: {original_text} → {text}")

    return text

def add_page_number(doc):
    """添加页码"""
    config = get_config()
    page_number_config = config.get('page_number', {})

    if not page_number_config.get('enabled', True):
        return

    try:
        # 获取文档的第一个节
        section = doc.sections[0]

        # 获取页脚
        footer = section.footer

        # 清空现有页脚内容
        for para in footer.paragraphs:
            para.clear()

        # 如果没有段落，添加一个
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]

        # 设置段落对齐方式
        position = page_number_config.get('position', 'center')
        if position == 'left':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif position == 'right':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 创建页码文本的XML
        from docx.oxml.shared import qn

        # 根据格式添加页码
        page_format = page_number_config.get('format', '1/x')
        if '1' in page_format:
            # 添加当前页码字段
            run = footer_para.add_run()
            fld_char_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_begin)
            instr_text = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
            run._r.append(instr_text)
            fld_char_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_end)

        if '/' in page_format:
            # 添加分隔符
            sep_run = footer_para.add_run("/")

        if 'x' in page_format:
            # 添加总页数字段
            total_run = footer_para.add_run()
            fld_char_begin2 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_begin2)
            instr_text2 = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> NUMPAGES </w:instrText>')
            total_run._r.append(instr_text2)
            fld_char_end2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_end2)

        # 设置字体格式
        font_name = page_number_config.get('font', 'Times New Roman')
        font_size = page_number_config.get('size', 10.5)

        for run in footer_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置西文字体
            run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

    except Exception as e:
        print(f"⚠️  页码添加失败，将跳过页码设置: {e}")
        pass

def parse_text_formatting(paragraph, text, title_level=0, is_quote=False):
    """解析文本格式（支持加粗、斜体、下划线，转换引号为中文）"""
    
    # 转换英文引号为中文引号
    text = convert_quotes_to_chinese(text)
    
    # 先处理<br>标签为段内换行
    segments = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    
    # 使用正则表达式解析所有格式标记
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]
    
    for idx, segment in enumerate(segments):
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = paragraph.add_run(part_text)
                set_run_format_with_styles(run, formats, title_level=title_level, is_quote=is_quote)
        if idx < len(segments) - 1:
            paragraph.add_run().add_break()

def parse_formatted_text(text, format_patterns):
    """解析带格式的文本，返回(文本, 格式)的列表"""
    
    if not text:
        return []
    
    parts = []
    current_pos = 0
    
    # 查找所有格式标记的位置
    all_matches = []
    for pattern, format_dict in format_patterns:
        for match in re.finditer(pattern, text):
            all_matches.append({
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'format': format_dict,
                'full_match': match.group(0)
            })
    
    # 按开始位置排序
    all_matches.sort(key=lambda x: x['start'])
    
    # 处理重叠的匹配（选择最长的匹配）
    filtered_matches = []
    for match in all_matches:
        # 检查是否与已有匹配重叠
        overlap = False
        for existing in filtered_matches:
            if (match['start'] < existing['end'] and match['end'] > existing['start']):
                # 有重叠，选择更长的匹配
                if len(match['full_match']) > len(existing['full_match']):
                    filtered_matches.remove(existing)
                    filtered_matches.append(match)
                overlap = True
                break
        if not overlap:
            filtered_matches.append(match)
    
    # 重新按位置排序
    filtered_matches.sort(key=lambda x: x['start'])
    
    # 构建文本部分列表
    for match in filtered_matches:
        # 添加前面的普通文本
        if current_pos < match['start']:
            normal_text = text[current_pos:match['start']]
            if normal_text:
                parts.append((normal_text, {}))
        
        # 添加格式化文本
        parts.append((match['text'], match['format']))
        current_pos = match['end']
    
    # 添加剩余的普通文本
    if current_pos < len(text):
        remaining_text = text[current_pos:]
        if remaining_text:
            parts.append((remaining_text, {}))
    
    # 如果没有找到任何格式，返回整个文本作为普通文本
    if not parts:
        parts.append((text, {}))
    
    return parts

def set_run_format(run, title_level=0):
    """设置文本运行格式（基础版本，用于标题）"""
    config = get_config()
    font_config = config.get('fonts.default', {})

    font = run.font
    font.name = font_config.get('ascii', 'Times New Roman')
    font.color.rgb = RGBColor(0, 0, 0)
    font.bold = False
    font.italic = False
    font.underline = False

    # 设置字体映射
    run._element.rPr.rFonts.set(qn('w:ascii'), font_config.get('ascii', 'Times New Roman'))
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_config.get('ascii', 'Times New Roman'))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('name', '仿宋_GB2312'))
    run._element.rPr.rFonts.set(qn('w:cs'), font_config.get('ascii', 'Times New Roman'))

    # 根据标题级别设置字号和加粗
    if title_level == 1:
        title_config = config.get('titles.level1', {})
        font.size = Pt(title_config.get('size', 15))
        font.bold = title_config.get('bold', True)
    elif title_level == 2:
        title_config = config.get('titles.level2', {})
        font.size = Pt(title_config.get('size', 12))
        font.bold = title_config.get('bold', True)
    elif title_level == 3:
        title_config = config.get('titles.level3', {})
        font.size = Pt(title_config.get('size', 12))
        font.bold = title_config.get('bold', False)
    elif title_level == 4:
        title_config = config.get('titles.level4', {})
        font.size = Pt(title_config.get('size', 12))
        font.bold = title_config.get('bold', False)
    else:
        font.size = Pt(font_config.get('size', 12))
        font.bold = False

def set_run_format_with_styles(run, formats, title_level=0, is_quote=False):
    """设置文本运行格式（支持多种样式）"""
    config = get_config()
    font_config = config.get('fonts.default', {})

    font = run.font
    font.name = font_config.get('ascii', 'Times New Roman')
    font.color.rgb = RGBColor(0, 0, 0)

    # 设置字体映射
    run._element.rPr.rFonts.set(qn('w:ascii'), font_config.get('ascii', 'Times New Roman'))
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_config.get('ascii', 'Times New Roman'))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('name', '仿宋_GB2312'))
    run._element.rPr.rFonts.set(qn('w:cs'), font_config.get('ascii', 'Times New Roman'))

    # 设置基础格式
    if title_level == 1:
        title_config = config.get('titles.level1', {})
        font.size = Pt(title_config.get('size', 15))
        font.bold = title_config.get('bold', True)
    elif title_level == 2:
        title_config = config.get('titles.level2', {})
        font.size = Pt(title_config.get('size', 12))
        font.bold = title_config.get('bold', True)
    elif is_quote:
        # 引用使用较小字号
        font.size = Pt(9)
        font.bold = False
    else:
        font.size = Pt(font_config.get('size', 12))
        font.bold = False

    # 应用Markdown格式
    if formats.get('code', False):
        code_config = config.get('inline_code', {})
        font.name = code_config.get('font', 'Times New Roman')
        font.size = Pt(code_config.get('size', 10))
        font.color.rgb = hex_to_rgb(code_config.get('color', '#333333'))
        run._element.rPr.rFonts.set(qn('w:ascii'), font.name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font.name)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font.name)
    elif formats.get('math', False):
        math_config = config.get('math', {})
        font.name = math_config.get('font', 'Times New Roman')
        font.size = Pt(math_config.get('size', 11))
        font.italic = math_config.get('italic', True)
        font.color.rgb = hex_to_rgb(math_config.get('color', '#00008B'))
        run._element.rPr.rFonts.set(qn('w:ascii'), font.name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font.name)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font.name)
    else:
        if formats.get('bold', False):
            font.bold = True
        if formats.get('italic', False):
            font.italic = True
        if formats.get('underline', False):
            font.underline = True
        if formats.get('strikethrough', False):
            font.strike = True

def add_horizontal_line(doc):
    """添加分割线"""
    config = get_config()
    hr_config = config.get('horizontal_rule', {})

    p = doc.add_paragraph()
    align_str = hr_config.get('alignment', 'center')
    p.alignment = parse_alignment(align_str)

    character = hr_config.get('character', '─')
    repeat_count = hr_config.get('repeat_count', 55)
    run = p.add_run(character * repeat_count)

    font_name = hr_config.get('font', 'Times New Roman')
    font_size = hr_config.get('size', 12)
    color_hex = hr_config.get('color', '#808080')

    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = hex_to_rgb(color_hex)

def add_task_list(doc, line):
    """添加任务列表"""
    config = get_config()
    task_config = config.get('lists.task', {})

    is_checked = line.startswith(('- [x]', '- [X]'))
    text = line[5:].strip()
    p = doc.add_paragraph()

    checked_mark = task_config.get('checked', '☑')
    unchecked_mark = task_config.get('unchecked', '☐')
    checkbox_run = p.add_run(f'{checked_mark} ' if is_checked else f'{unchecked_mark} ')
    set_run_format_with_styles(checkbox_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)

def add_bullet_list(doc, line):
    """添加无序列表"""
    config = get_config()
    bullet_config = config.get('lists.bullet', {})

    text = line[2:].strip()
    p = doc.add_paragraph()

    marker = bullet_config.get('marker', '•')
    bullet_run = p.add_run(f'{marker} ')
    set_run_format_with_styles(bullet_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_numbered_list(doc, line):
    """添加有序列表（保持原样输出）"""
    p = doc.add_paragraph()
    parse_text_formatting(p, line)
    set_paragraph_format(p)

def add_quote(doc, text):
    """添加带有着重底色的引用，并处理内部列表和多行文本"""
    config = get_config()
    quote_config = config.get('quote', {})

    # 按换行符分割文本，处理每一行
    lines = text.split('\n')

    # 获取配置
    bg_color = quote_config.get('background_color', '#EAEAEA')
    left_indent = quote_config.get('left_indent_inches', 0.2)
    font_size = quote_config.get('font_size', 9)
    line_spacing = quote_config.get('line_spacing', 1.5)

    for line_index, line in enumerate(lines):
        if not line.strip():  # 处理空行
            # 添加空段落来保持间距
            p = doc.add_paragraph()
            set_paragraph_format(p, is_quote=True)
            continue

        p = doc.add_paragraph()

        # 设置段落底色
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn

        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color.lstrip('#'))
        pPr.append(shd)

        # 设置左侧缩进
        p.paragraph_format.left_indent = Inches(left_indent)
        p.paragraph_format.line_spacing = line_spacing

        # 检查并处理列表标记
        bullet_match = re.match(r'^\s*([-*+])\s+', line)
        number_match = re.match(r'^\s*(\d+\.)\s+', line)

        list_marker_run = None

        if bullet_match:
            # 无序列表，使用符号并添加缩进
            indent_and_bullet = '    •  ' # 4个空格缩进 + bullet
            list_marker_run = p.add_run(indent_and_bullet)
            line = line[bullet_match.end():]
        elif number_match:
            # 有序列表，使用数字并添加缩进
            indent_and_number = f'    {number_match.group(1)} '
            list_marker_run = p.add_run(indent_and_number)
            line = line[number_match.end():]

        # 为列表标记设置统一格式
        if list_marker_run:
            list_marker_run.font.size = Pt(font_size)
            set_run_format_with_styles(list_marker_run, {}, is_quote=True)

        # 添加并解析文本内容
        parse_text_formatting(p, line, is_quote=True)

        # 调整段落格式和字号
        set_paragraph_format(p, is_quote=True)
        # 设置引用的字号
        for run in p.runs:
            run.font.size = Pt(font_size)

def add_code_block(doc, code_lines, language):
    """添加代码块"""
    config = get_config()
    code_config = config.get('code_block', {})

    # 语言标签配置
    label_config = code_config.get('label', {})
    if language:
        lang_p = doc.add_paragraph()
        lang_run = lang_p.add_run(f"[{language}]")
        lang_run.font.name = label_config.get('font', 'Times New Roman')
        lang_run.font.size = Pt(label_config.get('size', 10))
        lang_run.font.color.rgb = hex_to_rgb(label_config.get('color', '#808080'))

    # 代码内容配置
    content_config = code_config.get('content', {})
    left_indent = content_config.get('left_indent', 24)
    line_spacing = content_config.get('line_spacing', 1.2)
    font_name = content_config.get('font', 'Times New Roman')
    font_size = content_config.get('size', 10)
    color_hex = content_config.get('color', '#333333')

    for code_line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(code_line or ' ')
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = hex_to_rgb(color_hex)
        p.paragraph_format.left_indent = Pt(left_indent)
        p.paragraph_format.line_spacing = line_spacing

def set_paragraph_format(paragraph, title_level=0, is_quote=False):
    """设置段落格式"""
    config = get_config()
    paragraph_config = config.get('paragraph', {})

    # 设置段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = paragraph_config.get('line_spacing', 1.5)

    if title_level == 1:
        # 一级标题配置
        title_config = config.get('titles.level1', {})
        align_str = title_config.get('align', 'center')
        paragraph_format.alignment = parse_alignment(align_str)
        paragraph_format.space_before = Pt(title_config.get('space_before', 6))
        paragraph_format.space_after = Pt(title_config.get('space_after', 6))
        paragraph_format.first_line_indent = Pt(title_config.get('indent', 0))
    elif title_level == 2:
        # 二级标题配置
        title_config = config.get('titles.level2', {})
        align_str = title_config.get('align', 'justify')
        paragraph_format.alignment = parse_alignment(align_str)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(title_config.get('indent', 24))
    elif title_level == 3:
        # 三级标题配置
        title_config = config.get('titles.level3', {})
        align_str = title_config.get('align', 'justify')
        paragraph_format.alignment = parse_alignment(align_str)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(title_config.get('indent', 24))
    elif title_level == 4:
        # 四级标题配置
        title_config = config.get('titles.level4', {})
        align_str = title_config.get('align', 'justify')
        paragraph_format.alignment = parse_alignment(align_str)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(title_config.get('indent', 24))
    elif is_quote:
        # 引用：两端对齐，无首行缩进
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
    else:
        # 正文段落配置
        align_str = paragraph_config.get('align', 'justify')
        paragraph_format.alignment = parse_alignment(align_str)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(paragraph_config.get('first_line_indent', 24))

    # 确保所有runs都有正确的格式
    for run in paragraph.runs:
        if not hasattr(run.font, 'name') or not run.font.name:
            set_run_format(run, title_level)


def parse_alignment(align_str: str):
    """将字符串对齐方式转换为 WD_PARAGRAPH_ALIGNMENT 常量"""
    align_str = align_str.lower()
    if align_str == 'left':
        return WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align_str == 'center':
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align_str == 'right':
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:  # justify
        return WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def hex_to_rgb(hex_color: str):
    """将十六进制颜色转换为 RGBColor"""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 6:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    return RGBColor(0, 0, 0)  # 默认黑色

def find_template_file():
    """查找模板文件"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_dir = os.path.dirname(script_dir)  # 上级目录是 skill 根目录
    templates_dir = os.path.join(skill_dir, 'assets', 'templates')
    docx_files = glob.glob(os.path.join(templates_dir, "*.docx"))

    for docx_file in docx_files:
        filename = os.path.basename(docx_file).lower()
        if not any(keyword in filename for keyword in ['完整版', 'test', 'output', '输出']):
            if '模板' in filename or 'template' in filename:
                return docx_file

    return docx_files[0] if docx_files else None


def find_md_files():
    """查找脚本所在目录下的所有 .md 文件（与运行目录无关）"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_files = glob.glob(os.path.join(script_dir, "*.md"))
    return md_files

def generate_output_filename(md_file):
    """根据输入文件名生成输出文件名"""
    base_name = os.path.splitext(md_file)[0]
    return f"{base_name}_完整版.docx"

def main():
    """主函数"""

    # 创建参数解析器
    parser = argparse.ArgumentParser(
        description='Markdown到Word文档转换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.md                           # 使用默认配置
  %(prog)s input.md --preset=academic         # 使用学术论文预设
  %(prog)s input.md --config=my-config.yaml   # 使用自定义配置
  %(prog)s input.md output.docx               # 指定输出文件
  %(prog)s --list-presets                     # 列出所有预设
        """
    )

    parser.add_argument('input', nargs='?', help='输入的 Markdown 文件')
    parser.add_argument('output', nargs='?', help='输出的 Word 文件')
    parser.add_argument('--preset', '-p', help='使用预设配置 (legal/academic/report/simple)', default='legal')
    parser.add_argument('--config', '-c', help='使用自定义配置文件 (YAML格式)')
    parser.add_argument('--list-presets', action='store_true', help='列出所有可用的预设配置')
    parser.add_argument('--template', '-t', help='Word模板文件路径')

    args = parser.parse_args()

    # 处理 --list-presets
    if args.list_presets:
        print("可用的预设配置:")
        presets = list_presets()
        if presets:
            for preset in presets:
                config = get_preset(preset)
                if config:
                    print(f"  - {preset}: {config.description}")
        else:
            print("  没有可用的预设配置")
        return

    # 加载配置
    config = None
    if args.config:
        # 优先使用自定义配置
        config = load_config(args.config)
        if config is None:
            print(f"❌ 无法加载配置文件: {args.config}")
            return
        print(f"📋 使用配置文件: {args.config}")
    elif args.preset:
        # 使用预设配置
        config = get_preset(args.preset)
        if config is None:
            print(f"❌ 预设不存在: {args.preset}")
            print(f"可用预设: {', '.join(list_presets())}")
            return
        print(f"📋 使用预设: {args.preset} - {config.description}")

    if config is None:
        config = get_default_preset()

    # 设置全局配置
    set_config(config)

    # 如果没有指定输入文件，进入自动模式
    if not args.input:
        auto_mode(config)
        return

    # 处理单个文件
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = args.input
    if not os.path.isabs(md_file):
        alt = os.path.join(script_dir, md_file)
        if os.path.exists(alt):
            md_file = alt

    if not os.path.exists(md_file):
        print(f"❌ 错误: 找不到文件 {md_file}")
        return

    # 确定输出文件名
    if args.output:
        output_file = args.output
    else:
        output_file = generate_output_filename(md_file)

    # 确定模板文件
    template_file = args.template
    if not template_file:
        template_file = find_template_file()

    try:
        create_word_document(md_file, output_file, template_file, config)
        print_success_info(output_file, config)
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()


def auto_mode(config: Config):
    """自动模式：处理当前目录下的所有.md文件"""
    md_files = find_md_files()

    if not md_files:
        print("❌ 当前目录下没有找到.md文件")
        print("\n💡 使用方法:")
        print("1. 将此脚本放在包含.md文件的文件夹中")
        print("2. 或者运行: python md2word.py 输入文件.md")
        print("3. 或者运行: python md2word.py 输入文件.md 输出文件.docx")
        print("4. 或者运行: python md2word.py 输入文件.md --preset=academic")
        print("5. 或者运行: python md2word.py 输入文件.md --config=my-config.yaml")
        print("\n📋 可用预设:")
        presets = list_presets()
        if presets:
            for preset in presets:
                cfg = get_preset(preset)
                if cfg:
                    print(f"  - {preset}: {cfg.description}")
        return

    print(f"🔍 找到 {len(md_files)} 个Markdown文件:")
    for i, md_file in enumerate(md_files, 1):
        print(f"  {i}. {md_file}")

    print("\n开始转换...")

    # 查找模板文件
    template_file = find_template_file()

    success_count = 0
    for md_file in md_files:
        output_file = generate_output_filename(md_file)
        try:
            create_word_document(md_file, output_file, template_file, config)
            success_count += 1
        except Exception as e:
            print(f"❌ 处理 {md_file} 时出错: {e}")

    print(f"\n✅ 转换完成！成功处理 {success_count}/{len(md_files)} 个文件")
    print_success_info(None, config)

def print_success_info(filename=None, config: Config = None):
    """打印成功信息"""
    if config is None:
        config = get_config()

    print("\n📋 自动应用的格式:")

    # 页面设置
    page_config = config.get('page', {})
    print(f"📄 页面大小: {page_config.get('width', 21.0)}cm × {page_config.get('height', 29.7)}cm")
    print(f"📐 页边距: 上下{page_config.get('margin_top', 2.54)}cm，左右{page_config.get('margin_left', 3.18)}cm")

    # 字体设置
    font_config = config.get('fonts.default', {})
    print(f"📝 字体: {font_config.get('name', '仿宋_GB2312')}，{font_config.get('color', '#000000')}")

    # 字号设置
    print(f"📏 字号: {font_config.get('size', 12)}pt")

    # 段落设置
    paragraph_config = config.get('paragraph', {})
    print(f"📐 行距: {paragraph_config.get('line_spacing', 1.5)}倍")
    print(f"📝 首行缩进: {paragraph_config.get('first_line_indent', 24)}pt")

    # 标题设置
    title1_config = config.get('titles.level1', {})
    print(f"🎯 一级标题: {title1_config.get('size', 15)}pt，{'加粗' if title1_config.get('bold') else '常规'}，{title1_config.get('align', 'center')}")

    # 页码设置
    page_number_config = config.get('page_number', {})
    if page_number_config.get('enabled', True):
        print(f"📄 页码设置: {page_number_config.get('format', '1/x')}格式，{page_number_config.get('font', 'Times New Roman')}{page_number_config.get('size', 10.5)}pt")

    # 引号设置
    quotes_config = config.get('quotes', {})
    if quotes_config.get('convert_to_chinese', True):
        print("💬 引号转换: 英文引号自动转为中文引号")

    print("📊 表格支持: Markdown表格自动转换为Word表格，带边框格式")
    print("📈 图表支持: Mermaid图表本地渲染为高清图片插入Word文档")
    print("✨ 格式支持: 支持**加粗**、*斜体*、<u>下划线</u>、~~删除线~~格式")
    print("\n🎯 完全无需手动调整！直接可用！")

    if filename:
        print(f"\n📁 输出文件: {filename}")

def debug_quotes_in_file(file_path):
    """简化的引号调试"""
    print("🔍 检查文件中的引号...")
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 快速统计引号（区分ASCII与中文引号）
    ascii_double = content.count('"')
    chinese_open = content.count('“')
    chinese_close = content.count('”')

    print(f"📊 引号统计: ASCII双引号={ascii_double}, 中文开引号={chinese_open}, 中文闭引号={chinese_close}")

    # 只测试第一行包含引号的内容
    for i, line in enumerate(content.split('\n'), 1):
        if '"' in line:
            print(f"🎯 测试第{i}行: {line.strip()}")
            _ = convert_quotes_to_chinese(line.strip())
            break

    print("-" * 30)

if __name__ == "__main__":
    main() 

# 创建日期：250122 - 完整版本：支持表格转换和格式保留
